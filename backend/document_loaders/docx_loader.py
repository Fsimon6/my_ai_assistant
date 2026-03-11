"""
DOCX文档加载器
"""
from typing import Optional
import docx
from io import BytesIO
from .base_loader import DocumentLoader
from .models import Document, ProcessingConfig


class DOCXLoader(DocumentLoader):
    """DOCX加载器"""

    def __init__(self, config: Optional[ProcessingConfig] = None):
        super().__init__(config)
        self.supported_extensions = ['.docx', '.doc']

    def load(self, file_path: str) -> Document:
        """加载docx文件"""
        with open(file_path, 'rb') as f:
            file_bytes = f.read()

        return self.load_from_bytes(file_bytes, file_path)

    def load_from_bytes(self, file_bytes: bytes, filename: str) -> Document:
        """从字节加载DOCX"""
        metadata = self.extract_metadata(filename, file_bytes)
        try:
            # 使用python-docx读取DOCX
            doc = docx.Document(BytesIO(file_bytes))

            # 提取文本
            text_parts = []

            # 段落提取
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # 提取表格
            for table_idx, table in enumerate(doc.tables):
                table_text = f'\n[表格]{table_idx+1}\n'
                for row_idx, row in enumerate(table.rows):
                    row_text = '|'
                    for cell in row.cells:
                        row_text += f'{cell.text.strip()} |'
                    table_text += row_text + '\n'
                    text_parts.append(table_text)

            full_text = '\n'.join(text_parts)

            # 提取文档属性
            core_properties = doc.core_properties
            metadata.update({
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables),
                'author': str(core_properties.author) if core_properties.author else '',
                'title': str(core_properties.title) if core_properties.title else '',
                'subject': str(core_properties.subject) if core_properties.subject else '',
                'created': str(core_properties.created) if core_properties.created else '',
                'modified': str(core_properties.modified) if core_properties.modified else '',
                'last_modified_by': str(core_properties.last_modified_by) if core_properties.last_modified_by else '',
                'revision': str(core_properties.revision) if core_properties.revision else '',

            })

            return Document(
                content=full_text,
                metadata=metadata,
                source_type='docx'
            )

        except Exception as e:
            raise ValueError(f'DOCX解析失败：{str(e)}')

    def extract_styles(self, file_bytes: bytes) -> dict:
        """提取文档样式信息"""
        try:
            doc = docx.Document(BytesIO(file_bytes))
            styles_info = {}

            # 收集使用的样式
            for style in doc.styles:
                if style.type == 1:    # 段落样式
                    styles_info[style.name] = {
                        'type': 'paragraph',
                        'font_name': style.font.name if style.font.name else 0,
                        'font_size': style.font.size.pt if style.font.size else 0,
                        'bold': style.font.bold,
                        'italic': style.font.italic,
                    }

            # 统计样式使用情况
            style_usage = {}
            for paragraph in doc.paragraphs:
                style_name = paragraph.style.name
                style_usage[style_name] = style_usage.get(style_name, 0) + 1

            return {
                'styles': styles_info,
                'style_usage': style_usage
            }

        except Exception as e:
            print(f'样式提取失败：{e}')
            return {}
